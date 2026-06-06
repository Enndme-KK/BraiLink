from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.shortcuts import get_object_or_404
from django.conf import settings
from django.http import FileResponse
from django.utils import timezone
from django.db.models import Q
import requests
import os
import time
import random
from datetime import datetime, timedelta
import json
import threading
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re
from urllib.parse import quote
from .models import MedicalRecord, CTScan
from .serializers import MedicalRecordSerializer, MedicalRecordListSerializer, CTScanSerializer, CTScanCreateSerializer

SCAN_MODE_ALIASES = {
    '1': 't1',
    '2': 't2',
    '3': 't1ce',
    '4': 'flair',
    'plain': 't1',
    'enhanced': 't1ce',
    't1': 't1',
    't2': 't2',
    't1ce': 't1ce',
    'flair': 'flair',
}


def normalize_scan_mode(scan_mode):
    """Normalize legacy numeric scan modes to the TransUNet mode keys."""
    value = str(scan_mode or '').strip().lower()
    return SCAN_MODE_ALIASES.get(value, 't1')


def has_missing_processed_image(ct_scan):
    """Return True when a scan has an original image but no AI result image yet."""
    return bool(ct_scan.original_image) and not bool(ct_scan.processed_image)


# --- 新闻抓取与文件清理的通用工具 ---
import io
import chardet

def detect_html_encoding(response):
    """根据HTTP头、页面meta以及内容自动检测编码，避免强制UTF-8导致的乱码。"""
    # 1) 首先尝试 requests 的 apparent_encoding（需要 chardet/charset-normalizer 支持）
    enc = None
    try:
        enc = (response.encoding or '').lower() if response.encoding else None
        # 如果服务器标注为 utf-8/utf8/gbk/gb2312 等，直接用
        if enc in ('utf-8', 'utf8', 'gbk', 'gb2312'):
            return 'utf-8' if enc in ('utf-8', 'utf8') else enc
    except Exception:
        pass

    try:
        apparent = (getattr(response, 'apparent_encoding', None) or '').lower()
        if apparent in ('utf-8', 'utf8', 'gbk', 'gb2312'):
            return 'utf-8' if apparent in ('utf-8', 'utf8') else apparent
    except Exception:
        pass

    # 2) 从HTML <meta> 中探测
    try:
        head = response.content[:4096].decode('ascii', errors='ignore')
        m = re.search(r'charset\s*=\s*([\w-]+)', head, re.I)
        if m:
            meta_enc = m.group(1).lower()
            if meta_enc in ('utf-8', 'utf8', 'gbk', 'gb2312'):
                return 'utf-8' if meta_enc in ('utf-8', 'utf8') else meta_enc
    except Exception:
        pass

    # 3) 使用 chardet 对二进制内容进行猜测
    try:
        guess = chardet.detect(response.content or b'')
        enc = (guess.get('encoding') or '').lower()
        if enc:
            # 常见中文站点优先用 gbk/gb2312 或 utf-8
            if 'gb' in enc:
                return 'gbk'
            if 'utf' in enc:
                return 'utf-8'
            return enc
    except Exception:
        pass

    # 4) 兜底使用 utf-8
    return 'utf-8'

def clean_existing_news_files(news_text_dir):
    """
    清理新闻文本目录中的冗余和重复：
    - 删除历史的备份/临时扩展名文件（如 .bak/.refbak/.finalbak/.advbak 等）
    - 按标题去重，仅保留每个标题的一份 .txt 文件
    """
    if not os.path.isdir(news_text_dir):
        return

    # 先删除花样备份文件
    suffixes_to_remove = ('.bak', '.refbak', '.finalbak', '.advbak', '.decbak')
    removed = []
    kept = {}
    try:
        for name in os.listdir(news_text_dir):
            full = os.path.join(news_text_dir, name)
            if not os.path.isfile(full):
                continue
            # 删除非 .txt 的临时/备份文件
            if name.endswith(suffixes_to_remove) or ('.txt.' in name and not name.endswith('.txt')):
                try:
                    os.remove(full)
                    removed.append(name)
                except Exception:
                    pass
        # 再对 .txt 文件按标题去重
        title_map = {}
        for name in os.listdir(news_text_dir):
            if not name.lower().endswith('.txt'):
                continue
            full = os.path.join(news_text_dir, name)
            try:
                with open(full, 'r', encoding='utf-8', errors='ignore') as f:
                    first_line = f.readline().strip()
                # 期望格式为: 标题: xxx
                title = ''
                if first_line.startswith('标题:'):
                    title = first_line.split('标题:')[-1].strip()
                elif first_line.lower().startswith('title:'):
                    title = first_line.split(':', 1)[-1].strip()
                else:
                    # 无法识别标题，使用文件名去掉前缀编号作为标题
                    base = os.path.splitext(name)[0]
                    title = re.sub(r'^\d+_', '', base)

                norm = title.strip()
                if not norm:
                    norm = name

                if norm in title_map:
                    # 已有同标题，删除当前作为重复
                    try:
                        os.remove(full)
                        removed.append(name)
                    except Exception:
                        pass
                else:
                    title_map[norm] = name
                    kept[norm] = name
            except Exception:
                # 读取失败的也删除，避免脏文件残留
                try:
                    os.remove(full)
                    removed.append(name)
                except Exception:
                    pass
    finally:
        if removed:
            print(f"[新闻清理] 已删除备份/重复文件 {len(removed)} 个: {removed[:5]}{'...' if len(removed)>5 else ''}")
        print(f"[新闻清理] 当前有效新闻文件数量: {len(kept)}")

NEWS_CACHE_FILENAME = 'cached_news.json'
NEWS_CACHE_TTL = timedelta(days=1)

CATEGORY_KEYWORDS = {
    '科研': ['科研', '研究', '试验', '实验', '论文', '团队', '成果', '突破', '发现', '发表', '科研人员', '实验室'],
    '诊疗': ['诊疗', '治疗', '手术', '诊断', '临床', '病例', '医患', '问诊', '用药', '处方', '医生', '会诊', '医院发布', '复诊'],
    '政策': ['政策', '指南', '通知', '发布', '标准', '规定', '会议', '建设', '部署', '规划', '改革', '医保', '管理', '规范', '公示'],
    '健康': ['健康', '养生', '饮食', '生活', '保健', '运动', '习惯', '预防', '饮水', '睡眠', '心理', '防护', '体检', '康复']
}


def detect_category(title: str, summary: str) -> str:
    """根据标题和摘要推断新闻分类"""
    text = f"{title or ''} {summary or ''}".lower()
    for category, keywords in CATEGORY_KEYWORDS.items():
        for keyword in keywords:
            if keyword and keyword.lower() in text:
                return category
    return '健康'


def load_news_cache(cache_file_path):
    """Load cached news data from disk."""
    if not os.path.exists(cache_file_path):
        return None
    try:
        with open(cache_file_path, 'r', encoding='utf-8') as cache_file:
            data = json.load(cache_file)
            return data
    except Exception as err:
        print(f"[新闻缓存] 读取缓存失败: {err}")
        return None


def save_news_cache(cache_file_path, payload):
    """Persist news payload to disk."""
    try:
        os.makedirs(os.path.dirname(cache_file_path), exist_ok=True)
        with open(cache_file_path, 'w', encoding='utf-8') as cache_file:
            json.dump(payload, cache_file, ensure_ascii=False, indent=2)
        print(f"[新闻缓存] 已更新缓存: {cache_file_path}")
    except Exception as err:
        print(f"[新闻缓存] 保存缓存失败: {err}")


def should_refresh_cache(last_updated_str):
    """Determine if cached news needs refresh."""
    if not last_updated_str:
        return True
    try:
        last_updated = datetime.fromisoformat(last_updated_str)
        if timezone.is_naive(last_updated):
            last_updated = timezone.make_aware(last_updated, timezone=timezone.get_current_timezone())
    except Exception:
        return True
    return timezone.now() - last_updated >= NEWS_CACHE_TTL


def async_refresh_news(scrape_callable, cache_file_path):
    """Kick off asynchronous refresh of news cache."""
    def _runner():
        try:
            payload = scrape_callable()
            if payload and payload.get('news'):
                save_news_cache(cache_file_path, payload)
        except Exception as err:
            print(f"[新闻缓存] 异步刷新失败: {err}")

    threading.Thread(target=_runner, daemon=True).start()

class MedicalRecordViewSet(viewsets.ModelViewSet):
    """医疗记录管理视图集"""
    permission_classes = [IsAuthenticated]

    def _get_family_patient_ids(self, user):
        """Return patient ids bound to the current family user."""
        if user.user_type != 'family' or not hasattr(user, 'family_profile'):
            return []

        from families.models import FamilyPatientBinding

        return list(
            FamilyPatientBinding.objects
            .filter(family=user.family_profile)
            .values_list('patient_id', flat=True)
        )
    
    def get_queryset(self):
        user = self.request.user
        print(f"[get_queryset] 当前用户: {user.username}, 用户类型: {user.user_type}")

        # 默认预取病人/医生关联，避免 N+1
        base_qs = (
            MedicalRecord.objects
            .select_related('patient', 'patient__user', 'doctor', 'doctor__user')
            .prefetch_related('ct_scans')
        )

        if user.user_type == 'patient':
            queryset = base_qs.filter(patient__user=user).order_by('-visit_date', '-created_at')
            print(f"[get_queryset] 患者病历数量: {queryset.count()}")
            return queryset
        elif user.user_type == 'doctor':
            queryset = base_qs.all().order_by('-created_at')
            print(f"[get_queryset] 医生可见病历总数: {queryset.count()}")
            return queryset
        elif user.user_type == 'family':
            patient_ids = self._get_family_patient_ids(user)
            queryset = base_qs.filter(patient_id__in=patient_ids).order_by('-created_at')
            print(f"[get_queryset] 家属可见绑定患者病历数量: {queryset.count()}")
            return queryset

        print(f"[get_queryset] 无法识别的用户类型，返回空查询集")
        return MedicalRecord.objects.none()
    
    def get_serializer_class(self):
        if self.action == 'list':
            return MedicalRecordListSerializer
        return MedicalRecordSerializer

    def _ensure_record_access(self, request, medical_record):
        if request.user.user_type == 'patient':
            return medical_record.patient.user_id == request.user.id
        if request.user.user_type == 'doctor':
            # 医生：① 已接诊该病历 ② 该病历未分配（可接诊） ③ 通过挂号与患者建立过医患关系
            doctor_profile = getattr(request.user, 'doctor_profile', None)
            if not doctor_profile:
                return False
            if medical_record.doctor_id == doctor_profile.id:
                return True
            if medical_record.doctor_id is None:
                return True
            from appointments.models import Appointment
            return Appointment.objects.filter(
                doctor=doctor_profile,
                patient_id=medical_record.patient_id,
                status__in=['accepted', 'completed']
            ).exists()
        if request.user.user_type == 'family':
            return medical_record.patient_id in self._get_family_patient_ids(request.user)
        return False

    def _run_ai_analysis_for_ct_scan(self, ct_scan):
        """Run TransUNet inference for one existing CT/MRI scan and persist the result."""
        if not ct_scan.original_image:
            raise ValueError('该扫描没有原始影像，无法生成 AI 分割图')

        original_path = ct_scan.original_image.path
        if not os.path.exists(original_path):
            raise ValueError(f'原始影像文件不存在: {original_path}')

        from ml_service.prediction import (
            create_visualization,
            generate_segmentation_report,
            predict_segmentation,
            preprocess_mri_image,
        )

        scan_mode = normalize_scan_mode(ct_scan.scan_mode)
        output_dir = os.path.join(settings.MEDIA_ROOT, 'outputs')

        image_data = preprocess_mri_image(original_path, scan_mode)
        from ml_service.prediction import validate_medical_scan_image

        validation = validate_medical_scan_image(original_path)
        if not validation.get('is_medical_scan'):
            raise ValueError(validation.get('message') or '当前图片不像 CT/MRI 医学影像，已停止 AI 分析')

        pred_result = predict_segmentation(image_data, scan_mode)
        visualization_filename = create_visualization(
            image_data,
            pred_result['segmentation_mask'],
            scan_mode,
            output_dir
        )
        analysis = generate_segmentation_report(pred_result)

        processed_image_url = None
        if visualization_filename:
            processed_image_name = f'outputs/{visualization_filename}'
            ct_scan.processed_image.name = processed_image_name
            processed_image_url = f"{settings.MEDIA_URL}{processed_image_name}"

        ct_scan.scan_mode = scan_mode
        ct_scan.tumor_detected = pred_result.get('tumor_detected', False)
        ct_scan.tumor_type = 'unknown'
        ct_scan.tumor_size = str(pred_result.get('tumor_area', '')) if pred_result.get('tumor_area') else ''
        ct_scan.tumor_location = ''
        ct_scan.confidence_score = pred_result.get('confidence_score', 0.0)
        ct_scan.ai_analysis = analysis
        ct_scan.save()

        return {
            'ct_scan_id': ct_scan.id,
            'tumor_detected': ct_scan.tumor_detected,
            'tumor_area': pred_result.get('tumor_area', 0),
            'confidence_score': ct_scan.confidence_score,
            'scan_mode': pred_result.get('scan_mode', scan_mode),
            'analysis': analysis,
            'tumor_regions_count': len(pred_result.get('tumor_regions', [])),
            'processed_image': processed_image_url,
            'tumor_percentage': pred_result.get('tumor_percentage', 0),
            'class_statistics': pred_result.get('class_statistics', {}),
        }
    
    def destroy(self, request, *args, **kwargs):
        """删除医疗记录"""
        medical_record = self.get_object()

        # 权限检查：患者只能删除自己的病历；医生只能删除自己接诊或未分配的病历
        if request.user.user_type == 'patient':
            if medical_record.patient.user != request.user:
                return Response({
                    'error': '您只能删除自己的病历'
                }, status=status.HTTP_403_FORBIDDEN)
        elif request.user.user_type == 'doctor':
            if not self._ensure_record_access(request, medical_record):
                return Response({
                    'error': '只能删除自己接诊或未分配的病历'
                }, status=status.HTTP_403_FORBIDDEN)
        else:
            return Response({
                'error': '您没有权限删除病历'
            }, status=status.HTTP_403_FORBIDDEN)
        
        try:
            record_id = medical_record.id
            record_number = medical_record.record_number
            medical_record.delete()
            print(f"[删除病历] 用户 {request.user.username} 删除了病历 {record_id} ({record_number})")
            return Response({
                'message': '病历删除成功'
            }, status=status.HTTP_200_OK)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({
                'error': f'删除病历失败: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'])
    def complete_record(self, request, pk=None):
        """医生结案病历：将状态改为 completed 并通知患者。"""
        if request.user.user_type != 'doctor':
            return Response({'error': '只有医生可以结案病历'}, status=status.HTTP_403_FORBIDDEN)

        medical_record = self.get_object()
        if not self._ensure_record_access(request, medical_record):
            return Response({'error': '无权操作此病历'}, status=status.HTTP_403_FORBIDDEN)

        if medical_record.status == 'completed':
            return Response({'error': '该病历已结案'}, status=status.HTTP_400_BAD_REQUEST)

        medical_record.status = 'completed'
        medical_record.save(update_fields=['status'])

        # 通知患者
        try:
            from notifications.models import Notification
            Notification.objects.create(
                recipient=medical_record.patient.user,
                sender=request.user,
                notification_type='system',
                title='病历已结案',
                content=f'您的病历（{medical_record.record_number}）已由医生完成结案。',
            )
        except Exception:
            pass

        return Response({'message': '病历已结案', 'status': 'completed'})

    def create(self, request, *args, **kwargs):
        """创建医疗记录"""
        print(f"[创建病历] 用户类型: {request.user.user_type}")
        print(f"[创建病历] 接收到的数据: {request.data}")
        
        serializer = self.get_serializer(data=request.data)
        
        # 添加详细的验证错误信息
        if not serializer.is_valid():
            print(f"[创建病历] 验证失败: {serializer.errors}")
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # 根据用户类型设置不同的字段
        if request.user.user_type == 'doctor':
            # 医生创建病历
            print(f"[创建病历] 医生创建病历")
            medical_record = serializer.save(doctor=request.user.doctor_profile)
        elif request.user.user_type == 'patient':
            # 患者创建自己的病历（doctor字段可选）
            print(f"[创建病历] 患者创建病历")
            try:
                patient_profile = request.user.patient_profile
                print(f"[创建病历] 患者档案ID: {patient_profile.id}")
                
                # 强制使用当前患者的ID
                serializer.validated_data['patient'] = patient_profile
                
                # 如果提供了doctor_id，设置医生
                doctor_id = request.data.get('doctor_id')
                if doctor_id:
                    from doctors.models import Doctor
                    try:
                        doctor = Doctor.objects.get(id=doctor_id)
                        serializer.validated_data['doctor'] = doctor
                        if serializer.validated_data.get('status', 'pending') == 'pending':
                            serializer.validated_data['status'] = 'processing'
                        print(f"[创建病历] 指定医生: {doctor.name} (ID: {doctor.id})")
                    except Doctor.DoesNotExist:
                        return Response({'error': '指定的医生不存在'}, status=status.HTTP_400_BAD_REQUEST)
                
                medical_record = serializer.save()
                print(f"[创建病历] 病历创建成功，ID: {medical_record.id}, 医生: {medical_record.doctor}")
                
                # 如果指定了医生，创建挂号通知
                if medical_record.doctor:
                    try:
                        from notifications.models import Notification
                        from django.utils import timezone
                        
                        # 构建通知内容
                        visit_date_str = medical_record.visit_date.strftime('%Y-%m-%d %H:%M') if medical_record.visit_date else '待定'
                        department_str = medical_record.department or '未指定'
                        symptoms_str = medical_record.symptoms or '无'
                        
                        notification_content = f"""患者 {medical_record.patient.name} 已挂您的号，请及时接诊。

病历号：{medical_record.record_number}
就诊日期：{visit_date_str}
就诊科室：{department_str}
主诉：{symptoms_str}

请尽快处理该患者的挂号请求。"""
                        
                        # 创建通知
                        notification = Notification.objects.create(
                            recipient=medical_record.doctor.user,
                            sender=medical_record.patient.user,
                            notification_type='appointment',
                            title=f'新挂号通知 - {medical_record.patient.name}',
                            content=notification_content,
                            medical_record=medical_record
                        )
                        print(f"[创建通知] 已为医生 {medical_record.doctor.name} 创建挂号通知，通知ID: {notification.id}")
                    except Exception as e:
                        print(f"[创建通知] 创建挂号通知失败: {str(e)}")
                        # 通知创建失败不影响病历创建
                        
            except Exception as e:
                print(f"[创建病历] 创建失败: {str(e)}")
                import traceback
                traceback.print_exc()
                return Response({'error': f'创建病历失败: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response({'error': '无法识别的用户类型'}, status=status.HTTP_403_FORBIDDEN)
        
        return Response(MedicalRecordSerializer(medical_record).data, status=status.HTTP_201_CREATED)
    
    @action(detail=True, methods=['post'])
    def upload_ct_scan(self, request, pk=None):
        """上传CT扫描图片"""
        medical_record = self.get_object()
        
        # 允许医生上传任何病历的图片，患者只能上传自己的病历图片
        if request.user.user_type == 'patient':
            # 检查病历是否属于当前患者
            if medical_record.patient.user.id != request.user.id:
                return Response({'error': '您只能上传自己的医疗记录'}, status=status.HTTP_403_FORBIDDEN)
        elif request.user.user_type != 'doctor':
            return Response({'error': '您没有权限上传CT扫描'}, status=status.HTTP_403_FORBIDDEN)
        
        serializer = CTScanCreateSerializer(data=request.data, context={'medical_record': medical_record})
        if serializer.is_valid():
            ct_scan = serializer.save()
            print(f"[上传CT扫描] CT扫描记录已保存，ID: {ct_scan.id}")
            
            # 尝试使用已集成的 Django ML 服务进行AI分析（可选）
            ai_analysis_result = None
            try:
                print(f"[上传CT扫描] 尝试本地调用ML服务...")
                ai_analysis_result = self._run_ai_analysis_for_ct_scan(ct_scan)
                print(f"[上传CT扫描] AI分析成功，processed_image={ai_analysis_result.get('processed_image')}")
            except Exception as e:
                print(f"[上传CT扫描] AI分析出错: {str(e)}")
            
            # 无论AI分析是否成功，都返回成功响应（因为图片已经保存）
            response_data = {
                'ct_scan': CTScanSerializer(ct_scan).data,
                'message': 'CT扫描上传成功'
            }
            
            if ai_analysis_result:
                response_data['ai_analysis'] = ai_analysis_result
                response_data['message'] = 'CT扫描分析完成'
            else:
                response_data['warning'] = 'AI分析暂不可用，图片已保存'
            
            return Response(response_data, status=status.HTTP_201_CREATED)
        
        print(f"[上传CT扫描] 验证失败: {serializer.errors}")
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def reanalyze_ct_scans(self, request, pk=None):
        """为已有原图但缺少 AI 结果图的扫描补生成分割图。"""
        medical_record = self.get_object()

        if not self._ensure_record_access(request, medical_record):
            return Response({'error': '您没有权限重新分析该病历影像'}, status=status.HTTP_403_FORBIDDEN)

        force = bool(request.data.get('force', False))
        ct_scan_id = request.data.get('ct_scan_id')

        scans = medical_record.ct_scans.all().order_by('id')
        if ct_scan_id:
            scans = scans.filter(id=ct_scan_id)

        if not force:
            scans = scans.filter(Q(processed_image__isnull=True) | Q(processed_image=''))

        scans = [scan for scan in scans if scan.original_image]
        if not scans:
            return Response({
                'message': '没有需要补生成的 AI 分割图',
                'processed_count': 0,
                'errors': [],
                'ct_scans': CTScanSerializer(medical_record.ct_scans.all(), many=True).data,
            }, status=status.HTTP_200_OK)

        results = []
        errors = []
        for ct_scan in scans:
            try:
                print(f"[补生成AI分割] 开始处理 CTScan {ct_scan.id}")
                results.append(self._run_ai_analysis_for_ct_scan(ct_scan))
            except Exception as e:
                error = {'ct_scan_id': ct_scan.id, 'error': str(e)}
                errors.append(error)
                print(f"[补生成AI分割] CTScan {ct_scan.id} 失败: {str(e)}")

        status_code = status.HTTP_200_OK if results else status.HTTP_500_INTERNAL_SERVER_ERROR
        return Response({
            'message': f'已生成 {len(results)} 张 AI 分割图' if results else 'AI 分割图生成失败',
            'processed_count': len(results),
            'error_count': len(errors),
            'results': results,
            'errors': errors,
            'ct_scans': CTScanSerializer(medical_record.ct_scans.all(), many=True).data,
        }, status=status_code)
    
    @action(detail=True, methods=['get'])
    def ct_scans(self, request, pk=None):
        """获取医疗记录的CT扫描"""
        medical_record = self.get_object()
        ct_scans = medical_record.ct_scans.all()
        serializer = CTScanSerializer(ct_scans, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['put'])
    def verify_ct_scan(self, request, pk=None):
        """医生验证CT扫描结果"""
        medical_record = self.get_object()
        
        if request.user.user_type != 'doctor':
            return Response({'error': '只有医生可以验证结果'}, status=status.HTTP_403_FORBIDDEN)
        
        ct_scan_id = request.data.get('ct_scan_id')
        doctor_review = request.data.get('doctor_review', '')
        
        try:
            ct_scan = medical_record.ct_scans.get(id=ct_scan_id)
            ct_scan.doctor_review = doctor_review
            ct_scan.is_verified = True
            ct_scan.save()
            
            return Response({'message': '验证完成'})
        except CTScan.DoesNotExist:
            return Response({'error': 'CT扫描记录不存在'}, status=status.HTTP_404_NOT_FOUND)
    
    @action(detail=True, methods=['post'])
    def assign_to_me(self, request, pk=None):
        """医生接诊未分配的病历"""
        medical_record = self.get_object()
        
        if request.user.user_type != 'doctor':
            return Response({'error': '只有医生可以接诊'}, status=status.HTTP_403_FORBIDDEN)
        
        # 如果病历已经分配给其他医生，不允许接诊
        if medical_record.doctor is not None and medical_record.doctor.user.id != request.user.id:
            return Response({'error': '该病历已由其他医生接诊'}, status=status.HTTP_400_BAD_REQUEST)
        
        # 分配给当前医生
        medical_record.doctor = request.user.doctor_profile
        if medical_record.status == 'pending':
            medical_record.status = 'processing'
        medical_record.save(update_fields=['doctor', 'status', 'updated_at'])
        
        print(f"[接诊] 医生 {request.user.username} 接诊了病历 {medical_record.id}")
        
        # 更新相关的挂号通知为已读
        try:
            from notifications.models import Notification
            # 标记该医疗记录相关的挂号通知为已读
            Notification.objects.filter(
                medical_record=medical_record,
                notification_type='appointment',
                recipient=request.user,
                is_read=False
            ).update(is_read=True, read_at=timezone.now())
            print(f"[接诊] 已标记相关挂号通知为已读")
        except Exception as e:
            print(f"[接诊] 更新通知状态失败: {str(e)}")
            # 通知更新失败不影响接诊
        
        return Response({
            'message': '接诊成功',
            'medical_record': MedicalRecordSerializer(medical_record).data
        }, status=status.HTTP_200_OK)
    
    @action(detail=False, methods=['post'])
    def generate_word_report(self, request):
        """生成Word格式的医疗报告"""
        # 医生和患者都可以生成报告
        user_type = request.data.get('user_type', request.user.user_type)
        
        # 获取报告数据
        patient_data = request.data.get('patient', {})
        report_data = request.data.get('report', {})
        
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置文档样式
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            
            # 标题：医院名称
            title = doc.add_paragraph('天津市第一人民医院')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            
            # 副标题：报告类型
            subtitle = doc.add_paragraph('CT检查报告单')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.bold = True
            
            doc.add_paragraph()  # 空行
            
            # 患者信息部分
            doc.add_paragraph(f"病人号：{patient_data.get('pat_id', '未填写')}")
            doc.add_paragraph(f"检查日期：{patient_data.get('check_time', '未填写')}")
            doc.add_paragraph()
            
            info_para = doc.add_paragraph()
            info_para.add_run(f"姓名：{patient_data.get('name', '未填写')}").font.size = Pt(12)
            info_para.add_run(f"  性别：{patient_data.get('gender', '未填写')}").font.size = Pt(12)
            info_para.add_run(f"  年龄：{patient_data.get('age', '')}{'岁' if patient_data.get('age') else '未填写'}").font.size = Pt(12)
            
            doc.add_paragraph(f"住院号：{patient_data.get('host_id', '未填写')}")
            
            dept_para = doc.add_paragraph()
            dept_para.add_run(f"科室：{patient_data.get('department', '未填写')}").font.size = Pt(12)
            dept_para.add_run(f"  床号：{patient_data.get('bed_num', '未填写')}").font.size = Pt(12)
            
            doc.add_paragraph(f"报告日期：{patient_data.get('report_time', '未填写')}")
            check_project = patient_data.get('check_project', '未填写')
            if patient_data.get('position'):
                check_project += f" · {patient_data.get('position')}"
            doc.add_paragraph(f"检查项目：{check_project}")
            
            doc.add_paragraph()  # 空行
            
            # 影像结果部分
            finding_title = doc.add_paragraph('影像结果：')
            finding_title_run = finding_title.runs[0]
            finding_title_run.font.bold = True
            finding_title_run.font.size = Pt(14)
            
            finding_content = report_data.get('imaging_finding_result', '暂无检查结果')
            finding_para = doc.add_paragraph(finding_content)
            finding_para.paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进
            
            doc.add_paragraph()  # 空行
            
            # 影像诊断部分
            diagnosis_title = doc.add_paragraph('影像诊断：')
            diagnosis_title_run = diagnosis_title.runs[0]
            diagnosis_title_run.font.bold = True
            diagnosis_title_run.font.size = Pt(14)
            diagnosis_title_run.font.color.rgb = RGBColor(135, 77, 0)  # 橙色
            
            diagnosis_content = report_data.get('imaging_diagnosis_result', '暂无诊断意见')
            diagnosis_para = doc.add_paragraph()
            diagnosis_para.paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进
            # 诊断内容使用稍微深一点的橙色
            diagnosis_run = diagnosis_para.add_run(diagnosis_content)
            diagnosis_run.font.color.rgb = RGBColor(135, 77, 0)
            
            doc.add_paragraph()  # 空行
            doc.add_paragraph('─' * 50)  # 分隔线
            doc.add_paragraph()  # 空行
            
            # 医生签名部分
            doctor_info = doc.add_paragraph()
            doctor_info.add_run(f"报告医师：{report_data.get('doctor', ['', ''])[0] or '待填写'}").font.size = Pt(12)
            doctor_info.add_run(f"   审核医师：{report_data.get('doctor', ['', ''])[1] or '待填写'}").font.size = Pt(12)
            
            # 保存文件到统一 output/reports 目录
            user_type_folder = user_type if user_type in ('doctor', 'patient', 'family') else 'patient'
            output_dir = os.path.join(settings.REPORT_OUTPUT_ROOT, user_type_folder)
            
            os.makedirs(output_dir, exist_ok=True)
            
            # 生成文件名
            patient_name = patient_data.get('name', '患者') or '患者'
            check_time = patient_data.get('check_time', '').replace(':', '-').replace(' ', '_') or datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_name = "".join(c for c in patient_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"医疗报告_{safe_name}_{check_time}.docx"
            filepath = os.path.join(output_dir, filename)
            
            # 保存文档
            doc.save(filepath)
            
            # 检查文件是否成功保存
            if not os.path.exists(filepath):
                raise Exception(f'文件保存失败: {filepath}')
            
            # 返回文件路径和下载URL（文件在output文件夹，需要通过其他方式访问）
            # 返回绝对路径用于确认文件位置
            absolute_path = os.path.abspath(filepath)
            return Response({
                'message': '报告生成成功',
                'file_path': absolute_path,  # 返回绝对路径
                'relative_path': f'output/reports/{user_type_folder}/{filename}',  # 相对路径
                'filename': filename,
                'output_dir': output_dir,
                'user_type': user_type
            }, status=status.HTTP_201_CREATED)
            
        except Exception as e:
            return Response({
                'error': f'生成报告失败: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_medical_news(request):
    """获取新浪网医学新闻"""
    stale_cache = None
    try:
        # 创建保存目录
        backend_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        news_dir = os.path.join(backend_dir, 'media', 'news')
        news_text_dir = os.path.join(news_dir, 'texts')
        news_image_dir = os.path.join(news_dir, 'images')
        os.makedirs(news_text_dir, exist_ok=True)
        os.makedirs(news_image_dir, exist_ok=True)
        
        # 先清理历史备份与重复新闻文件，避免前端读到乱码或重复
        try:
            clean_existing_news_files(news_text_dir)
        except Exception as _e:
            print(f"[新闻清理] 执行清理时出现问题: {_e}")
        
        # 读取缓存（每日刷新一次）
        cache_file_path = os.path.join(news_dir, NEWS_CACHE_FILENAME)
        cache_payload = load_news_cache(cache_file_path)
        force_refresh = request.query_params.get('refresh', '').lower() in ('1', 'true', 'yes')
        
        existing_news = []
        existing_titles = set()
        if cache_payload and cache_payload.get('news'):
            for item in cache_payload.get('news', []):
                title = (item.get('title') or '').strip()
                if not title:
                    continue
                summary = item.get('summary') or ''
                if not item.get('category'):
                    item['category'] = detect_category(title, summary)
                existing_news.append(item)
                existing_titles.add(title)
        
        if existing_news and not force_refresh:
            cache_is_stale = should_refresh_cache(cache_payload.get('timestamp'))
            if not cache_is_stale:
                cache_payload['from_cache'] = True
                cache_payload['cache_stale'] = False
                return Response(cache_payload)
            else:
                print("[新闻缓存] 缓存超过TTL，准备刷新")
                stale_cache = cache_payload
        elif cache_payload:
            stale_cache = cache_payload
        
        # 随机User-Agent列表，模拟不同浏览器
        user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:121.0) Gecko/20100101 Firefox/121.0',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.1 Safari/605.1.15',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        ]
        
        # 设置请求头，模拟真实浏览器
        headers = {
            'User-Agent': random.choice(user_agents),
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Cache-Control': 'max-age=0',
            'Referer': 'https://www.sina.com.cn/'
        }
        
        # 请求新浪网医学健康频道（更权威的医学新闻源）
        # 尝试多个医学健康相关的URL
        medical_urls = [
            'https://health.sina.com.cn/',  # 新浪健康频道
            'https://med.sina.com/',  # 新浪医学频道
            'https://www.sina.com.cn/',  # 新浪首页（备用）
        ]
        
        # 使用Session保持连接
        session = requests.Session()
        session.headers.update(headers)
        
        soups = []
        max_retries = 3

        def fetch_soup(target_url):
            nonlocal max_retries
            response_local = None
            for retry in range(max_retries):
                try:
                    response_local = session.get(target_url, timeout=(10, 30), allow_redirects=True)
                    response_local.encoding = 'utf-8'
                    if response_local.status_code == 200:
                        break
                    else:
                        print(f"[获取新闻] 请求失败，状态码: {response_local.status_code}，重试 {retry + 1}/{max_retries}")
                        if retry < max_retries - 1:
                            time.sleep(random.uniform(2.0, 4.0))
                except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
                    print(f"[获取新闻] 请求超时或连接错误: {str(e)}，重试 {retry + 1}/{max_retries}")
                    if retry < max_retries - 1:
                        time.sleep(random.uniform(3.0, 5.0))
                except Exception as e:
                    print(f"[获取新闻] 请求异常: {str(e)}")
                    return None
            if not response_local or response_local.status_code != 200:
                print(f"[获取新闻] 请求失败，状态码: {response_local.status_code if response_local else 'None'}")
                return None
            try:
                enc = detect_html_encoding(response_local)
                html = response_local.content.decode(enc, errors='replace')
                return BeautifulSoup(html, 'lxml')
            except Exception as err:
                print(f"[获取新闻] HTML解析失败: {str(err)}")
                try:
                    return BeautifulSoup(html, 'html.parser')
                except Exception:
                    return None

        for index_url, target_url in enumerate(medical_urls):
            time.sleep(random.uniform(1.5, 3.0))
            soup = fetch_soup(target_url)
            if soup is not None:
                soups.append(soup)
            if len(soups) >= 3:  # 已经抓取主要的三个入口
                break

        if not soups:
            print("[获取新闻] 所有入口均获取失败")
            return Response({
                'news': []
            })
        
        news_list = []
        seen_titles = set(existing_titles)  # 用于去重，包含历史数据
        MAX_CANDIDATES = 120
        TARGET_NEWS_COUNT = 60
        aggregated_high_quality = []
        aggregated_normal = []
        collected_titles = set(existing_titles)
        
        def has_existing_title(title_text):
            normalized_title = title_text.strip()
            if not normalized_title:
                return True
            if normalized_title in collected_titles:
                return True
            return any(existing[1] == normalized_title for existing in aggregated_high_quality) or \
                   any(existing[1] == normalized_title for existing in aggregated_normal)
        
        def add_candidate(link_obj, title_text, href_value, is_high_quality):
            normalized_title = title_text.strip()
            if has_existing_title(normalized_title):
                return False
            candidate = (link_obj, normalized_title, href_value, is_high_quality)
            if is_high_quality:
                aggregated_high_quality.append(candidate)
            else:
                aggregated_normal.append(candidate)
            collected_titles.add(normalized_title)
            total = len(aggregated_high_quality) + len(aggregated_normal)
            return total >= MAX_CANDIDATES
        
        # 辅助函数：处理图片URL（相对路径转绝对路径）
        def process_image_url(img_src):
            """处理图片URL，转换为完整的绝对URL"""
            if not img_src:
                return None
            if img_src.startswith('http'):
                return img_src
            elif img_src.startswith('//'):
                return 'https:' + img_src
            elif img_src.startswith('/'):
                return 'https://www.sina.com.cn' + img_src
            else:
                return 'https://www.sina.com.cn/' + img_src
        
        # 辅助函数：下载并保存图片
        def download_and_save_image(image_url, news_title, index):
            """下载图片并保存到本地"""
            try:
                if not image_url or image_url.startswith('data:') or 'placeholder' in image_url:
                    return None
                
                # 获取图片扩展名
                ext = 'jpg'
                if '.jpg' in image_url.lower() or '.jpeg' in image_url.lower():
                    ext = 'jpg'
                elif '.png' in image_url.lower():
                    ext = 'png'
                elif '.gif' in image_url.lower():
                    ext = 'gif'
                elif '.webp' in image_url.lower():
                    ext = 'webp'
                
                # 生成文件名（使用hash避免中文文件名问题）
                import hashlib
                # 使用标题和索引生成hash，确保唯一性
                title_hash = hashlib.md5(f"{index}_{news_title}".encode('utf-8')).hexdigest()[:12]
                
                # 同时保留部分可读的标题（只保留ASCII字符）
                readable_title = "".join(c for c in news_title[:20] if ord(c) < 128 and (c.isalnum() or c in ('-', '_'))).strip()
                if not readable_title:
                    readable_title = "news"
                
                filename = f"{index}_{readable_title}_{title_hash}.{ext}"
                filepath = os.path.join(news_image_dir, filename)
                
                # 下载图片
                img_response = session.get(image_url, timeout=10, stream=True)
                if img_response.status_code == 200:
                    with open(filepath, 'wb') as f:
                        for chunk in img_response.iter_content(1024):
                            f.write(chunk)
                    print(f"[保存图片] 已保存: {filename}")
                    
                    # 返回相对路径（文件名已不包含中文，无需编码）
                    return f"/media/news/images/{filename}"
                else:
                    print(f"[保存图片] 下载失败，状态码: {img_response.status_code}")
                    return None
            except Exception as e:
                print(f"[保存图片] 保存图片失败: {str(e)}")
                return None
        
        # 辅助函数：保存新闻文本
        def save_news_text(news_title, news_content, news_url, index):
            """保存新闻文本到文件"""
            try:
                safe_title = "".join(c for c in news_title[:50] if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_title = safe_title.replace(' ', '_')
                filename = f"{index}_{safe_title}.txt"
                filepath = os.path.join(news_text_dir, filename)
                
                # 确保内容编码正确
                if isinstance(news_content, bytes):
                    news_content = news_content.decode('utf-8', errors='ignore')
                elif not isinstance(news_content, str):
                    news_content = str(news_content)
                
                # 清理内容，移除重复和无效字符
                lines = news_content.split('\n')
                cleaned_lines = []
                seen_lines = set()
                for line in lines:
                    line = line.strip()
                    if line and len(line) > 5 and line not in seen_lines:
                        # 过滤掉明显的导航文本
                        if not any(keyword in line for keyword in ['新浪', 'SINA', '首页', '导航', '菜单', '登录', '注册']):
                            cleaned_lines.append(line)
                            seen_lines.add(line)
                
                cleaned_content = '\n'.join(cleaned_lines)
                
                with open(filepath, 'w', encoding='utf-8', errors='ignore') as f:
                    f.write(f"标题: {news_title}\n")
                    f.write(f"链接: {news_url}\n")
                    f.write(f"日期: {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"{'='*50}\n\n")
                    f.write(f"内容:\n{cleaned_content}\n")
                
                print(f"[保存文本] 已保存: {filename}")
            except Exception as e:
                print(f"[保存文本] 保存文本失败: {str(e)}")
                import traceback
                traceback.print_exc()
        
        # 权威医学关键词（更专业、更有含金量）
        # 优先选择：疾病、治疗、研究、发现、突破、指南、共识、临床等关键词
        authoritative_keywords = [
            '疾病', '治疗', '研究', '发现', '突破', '指南', '共识', '临床', '诊断', '手术',
            '药物', '疫苗', '癌症', '肿瘤', '心血管', '糖尿病', '高血压', '心脏病', '脑卒中',
            '医学', '医疗', '医院', '医生', '患者', '康复', '预防', '筛查', '检测'
        ]
        
        # 高质量医学新闻特征关键词
        quality_keywords = [
            '研究', '发现', '突破', '新进展', '指南', '共识', '临床', '病例', '治疗',
            '诊断', '手术', '药物', '疫苗', '预防', '筛查', '检测', '康复'
        ]
        total_break = False
        for soup in soups:
            # 查找新闻链接（优先查找权威医学新闻）
            all_links = soup.find_all('a', href=True)
            
            for link in all_links:
                link_text = link.get_text(strip=True)
                href = link.get('href', '')
                
                if not link_text or len(link_text) < 10 or len(link_text) > 200:
                    continue
                
                # 检查是否包含医学关键词
                has_medical_keyword = any(keyword in link_text for keyword in authoritative_keywords)
                has_quality_keyword = any(keyword in link_text for keyword in quality_keywords)
                
                if has_medical_keyword and href and (href.startswith('http') or href.startswith('/')):
                    # 检查URL是否指向医学健康相关页面
                    href_lower = href.lower()
                    is_medical_url = any(keyword in href_lower for keyword in ['health', 'medical', 'med', 'yiliao', 'jiankang', 'disease'])
                    reached_limit = add_candidate(link, link_text, href, has_quality_keyword or is_medical_url)
                    if reached_limit:
                        total_break = True
                        break
            if total_break:
                break
            
            # 如果链接不够，从新闻区域补充
            if len(aggregated_high_quality) + len(aggregated_normal) < 40:
                news_containers = soup.find_all(['div', 'ul', 'li', 'article'], 
                                               class_=re.compile(r'news|article|item|list|content|main', re.I))
                
                for container in news_containers:
                    links = container.find_all('a', href=True)
                    for link in links:
                        link_text = link.get_text(strip=True)
                        href = link.get('href', '')
                        
                        if link_text and len(link_text) > 10 and len(link_text) < 200:
                            if href and (href.startswith('http') or href.startswith('/')):
                                has_quality = any(keyword in link_text for keyword in quality_keywords)
                                if add_candidate(link, link_text, href, has_quality):
                                    total_break = True
                                    break
                    if total_break or len(aggregated_high_quality) + len(aggregated_normal) >= 40:
                        break
            if total_break:
                break
        
        all_candidate_links = aggregated_high_quality + aggregated_normal
        
        print(f"[获取新闻] 找到 {len(all_candidate_links)} 个候选链接（高质量: {len(aggregated_high_quality)}）")
        
        # 处理每个链接，获取详细信息（优先处理高质量新闻）
        # 按质量排序：高质量在前
        all_candidate_links.sort(key=lambda x: x[3], reverse=True)  # 按质量排序
        
        # 处理更多候选链接，确保能获取足够的新新闻
        for idx, (link, title_text, href, is_high_quality) in enumerate(all_candidate_links[:MAX_CANDIDATES]):  # 最多处理候选
            if len(news_list) >= TARGET_NEWS_COUNT:  # 获取更多有图片的新闻
                break
            
            try:
                # 处理URL
                if not href.startswith('http'):
                    if href.startswith('/'):
                        link_url = 'https://www.sina.com.cn' + href
                    else:
                        link_url = 'https://www.sina.com.cn/' + href
                else:
                    link_url = href
                
                # 过滤无效标题
                if not title_text or len(title_text) < 10 or len(title_text) > 200:
                    continue
                
                # 去重
                if title_text in seen_titles:
                    continue
                seen_titles.add(title_text)
                
                # 查找图片（优先从详情页获取，避免获取到新浪logo等）
                image_url = None
                
                # 辅助函数：检查图片是否有效（不是logo、图标等）
                def is_valid_image(img_url, img_element=None):
                    """检查图片URL是否有效（不是logo、图标、水印等）"""
                    if not img_url:
                        return False
                    
                    url_lower = img_url.lower()
                    
                    # 过滤无效图片模式
                    invalid_patterns = [
                        'logo', 'icon', 'avatar', 'ad', 'advertisement', 
                        'watermark', 'sina_logo', 'sina_icon', 'sina.com.cn/logo',
                        'nav', 'menu', 'header', 'footer', 'sidebar',
                        'button', 'btn', 'arrow', 'close', 'share',
                        'thumb', 'thumbnail', 'small', 'tiny'
                    ]
                    
                    # 检查URL
                    for pattern in invalid_patterns:
                        if pattern in url_lower:
                            return False
                    
                    # 检查图片元素
                    if img_element:
                        img_class = img_element.get('class', [])
                        img_id = img_element.get('id', '')
                        if any(pattern in str(img_class).lower() or pattern in str(img_id).lower() 
                               for pattern in ['logo', 'icon', 'ad', 'nav', 'menu']):
                            return False
                        
                        # 检查尺寸
                        img_width = img_element.get('width')
                        img_height = img_element.get('height')
                        if img_width and img_height:
                            try:
                                w, h = int(img_width), int(img_height)
                                if w < 200 or h < 200:
                                    return False
                            except:
                                pass
                    
                    return True
                
                # 方法1: 从链接元素附近查找图片（但优先等待详情页）
                # 这里先不设置，等访问详情页后再决定
                
                # 方法3: 访问详情页获取图片和内容
                news_content = ''
                if link_url and link_url.startswith('http'):
                    try:
                        time.sleep(random.uniform(1.0, 2.0))
                        detail_response = session.get(link_url, headers=headers, timeout=(10, 30), allow_redirects=True)
                        if detail_response.status_code == 200:
                            # 自动探测编码，避免强制UTF-8导致的乱码
                            try:
                                denc = detect_html_encoding(detail_response)
                                dhtml = detail_response.content.decode(denc, errors='replace')
                            except Exception:
                                dhtml = detail_response.text
                            
                            try:
                                detail_soup = BeautifulSoup(dhtml, 'html.parser')
                            except Exception:
                                try:
                                    detail_soup = BeautifulSoup(dhtml, 'lxml')
                                except Exception:
                                    detail_soup = None
                            
                            # 获取文章内容区域（优先选择文章主体）
                            content_selectors = [
                                detail_soup.find(['article', 'div'], class_=re.compile(r'content|article|main|post|detail|body|text|art_content|article_content', re.I)),
                                detail_soup.find('div', id=re.compile(r'content|article|main|post|detail|body|art_content|article_content', re.I)),
                                detail_soup.find('main'),
                                detail_soup.find('article'),
                            ]
                            
                            # 提取文本内容
                            for content_area in content_selectors:
                                if content_area:
                                    # 排除导航、侧边栏等非内容区域
                                    if content_area.find_parent(class_=re.compile(r'nav|sidebar|menu|header|footer', re.I)):
                                        continue
                                    
                                    # 提取文本内容，确保编码正确
                                    paragraphs = content_area.find_all(['p', 'div'], limit=50)
                                    text_parts = []
                                    for p in paragraphs:
                                        text = p.get_text(strip=True)
                                        if text and len(text) > 10:  # 过滤太短的文本
                                            text_parts.append(text)
                                    
                                    if text_parts:
                                        news_content = '\n'.join(text_parts)
                                        if len(news_content) > 50:  # 确保内容足够长
                                            break
                            
                            # 优先从文章内容区域获取图片（而不是导航栏、侧边栏等）
                            if not image_url:
                                for content_area in content_selectors:
                                    if not content_area:
                                        continue
                                    
                                    # 排除导航、侧边栏等非内容区域
                                    if content_area.find_parent(class_=re.compile(r'nav|sidebar|menu|header|footer|ad|advertisement', re.I)):
                                        continue
                                    
                                    # 查找文章内容中的图片
                                    for img in content_area.find_all('img', limit=10):
                                        img_src = (img.get('src') or 
                                                  img.get('data-src') or 
                                                  img.get('data-original') or
                                                  img.get('data-lazy-src') or
                                                  img.get('data-url'))
                                        
                                        if not img_src:
                                            continue
                                        
                                        processed_url = process_image_url(img_src)
                                        if not processed_url:
                                            continue
                                        
                                        # 使用验证函数检查图片是否有效
                                        if is_valid_image(processed_url, img):
                                            image_url = processed_url
                                            print(f"[获取新闻] 从文章内容获取图片: {image_url[:80]}")
                                            break
                                    
                                    if image_url:
                                        break
                    except Exception as e:
                        print(f"[获取新闻] 访问详情页失败: {str(e)}")
                        import traceback
                        traceback.print_exc()
                
                # 如果没有内容，使用标题作为内容
                if not news_content:
                    news_content = title_text
                
                # 下载并保存图片
                saved_image_path = None
                if image_url:
                    saved_image_path = download_and_save_image(image_url, title_text, len(news_list) + 1)
                
                # 保存新闻文本
                save_news_text(title_text, news_content, link_url, len(news_list) + 1)
                
                # 如果图片保存成功，使用本地路径；否则使用原始URL
                # 如果没有图片，返回None，让前端处理
                final_image_url = saved_image_path if saved_image_path else (image_url if image_url else None)
                
                # 只添加有图片的新闻（确保轮播效果）
                if final_image_url:
                    summary_text = news_content[:100] + '...' if len(news_content) > 100 else news_content
                    news_list.append({
                        'title': title_text,
                        'summary': summary_text,
                        'url': link_url,
                        'date': timezone.now().strftime('%Y-%m-%d'),
                        'image': final_image_url,
                        'quality': 'high' if is_high_quality else 'normal',
                        'category': detect_category(title_text, summary_text)
                    })
                    
                    print(f"[获取新闻] 成功获取{'[高质量]' if is_high_quality else ''}: {title_text[:50]}, 图片: {final_image_url[:60] if final_image_url else '无'}")
                else:
                    print(f"[获取新闻] 跳过无图片新闻: {title_text[:50]}")
                
            except Exception as e:
                print(f"[获取新闻] 处理新闻项时出错: {str(e)}")
                continue
        
        # 返回新闻列表（包含标题和图片）
        # 按质量排序：高质量新闻在前
        news_list.sort(key=lambda x: (x.get('quality', 'normal') == 'high', x.get('title', '')), reverse=True)
        
        print(f"[获取新闻] 成功获取 {len(news_list)} 条新闻（带图片），返回前 {min(TARGET_NEWS_COUNT, len(news_list))} 条")
        
        # 将新新闻与历史缓存合并，保留唯一标题
        combined_news = news_list + existing_news
        deduped_news = []
        combined_seen = set()
        for item in combined_news:
            title = (item.get('title') or '').strip()
            if not title or title in combined_seen:
                continue
            combined_seen.add(title)
            summary = item.get('summary') or ''
            item['category'] = item.get('category') or detect_category(title, summary)
            deduped_news.append(item)
        
        # 返回更多新闻，让前端可以选择显示最新的
        payload = {
            'news': deduped_news[:TARGET_NEWS_COUNT],  # 保留最多60条供前端分类
            'total': len(deduped_news),  # 返回总数
            'timestamp': timezone.now().isoformat()  # 返回时间戳，用于判断是否刷新
        }
        save_news_cache(cache_file_path, payload)
        payload['from_cache'] = False
        payload['cache_stale'] = False
        return Response(payload)
        
    except Exception as e:
        print(f"[获取新闻] 错误: {str(e)}")
        import traceback
        traceback.print_exc()
        # 如果存在旧缓存，返回旧数据以保证页面可用
        if stale_cache and stale_cache.get('news'):
            print("[新闻缓存] 使用旧缓存作为降级数据")
            stale_cache['from_cache'] = True
            stale_cache['cache_stale'] = True
            return Response(stale_cache)
        # 返回空列表，让前端处理空状态
        print(f"[获取新闻] 发生异常，返回空列表")
        return Response({
            'news': [],  # 返回空列表，让前端处理空状态
            'total': 0,
            'timestamp': timezone.now().isoformat(),
            'from_cache': True,
            'cache_stale': True
        })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_today_patients(request):
    """获取今日就诊患者列表（医生端）"""
    if request.user.user_type != 'doctor':
        return Response({'error': '只有医生可以查看今日就诊患者'}, status=status.HTTP_403_FORBIDDEN)
    
    try:
        # 获取今天的日期范围
        today = timezone.now().date()
        today_start = timezone.make_aware(datetime.combine(today, datetime.min.time()))
        today_end = timezone.make_aware(datetime.combine(today, datetime.max.time()))
        
        # 查询今日的医疗记录
        # 包括：1. 今日创建的记录 2. 今日有更新的记录 3. 未分配的病历（患者自己添加的）
        today_records = MedicalRecord.objects.filter(
            Q(created_at__date=today) | 
            Q(updated_at__date=today) |
            Q(visit_date__date=today)
        ).order_by('-created_at')
        
        # 统计信息
        total_count = today_records.count()
        unassigned_count = today_records.filter(doctor__isnull=True).count()
        my_patients_count = today_records.filter(doctor__user=request.user).count()
        
        # 序列化数据，传递request上下文以便判断是否当前医生接诊
        serializer = MedicalRecordListSerializer(today_records, many=True, context={'request': request})
        
        return Response({
            'total_count': total_count,
            'unassigned_count': unassigned_count,
            'my_patients_count': my_patients_count,
            'patients': serializer.data
        })
        
    except Exception as e:
        print(f"[获取今日患者] 错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({
            'error': f'获取今日患者失败: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
